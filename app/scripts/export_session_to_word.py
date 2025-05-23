"""
Export the current session (metadata + exercise/music list) to a MS Word (.docx) file.
"""
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from datetime import datetime
import sys

# Add project root to sys.path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))

from app.db.queries import get_all_songs
from app.ui.components import get_song_file_path


def export_session_to_word(session_metadata, session_exercises, export_path=None):
    """
    Export the session metadata and exercise/music list to a Word document.
    """
    load_dotenv()
    # Convert session_metadata to a plain dict if it's a sqlite3.Row
    if hasattr(session_metadata, 'keys') and not isinstance(session_metadata, dict):
        session_metadata = {k: session_metadata[k] for k in session_metadata.keys()}
    if export_path is None:
        export_path = os.getenv("EXPORT_PATH", os.getcwd())
    if not os.path.exists(export_path):
        os.makedirs(export_path)
    session_name = session_metadata.get("name", "Session")
    filename = f"{session_name}_music_list.docx"
    filepath = os.path.join(export_path, filename)

    doc = Document()
    # Set page orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height    
    # Set margins
    doc.add_heading(f"Session: {session_name}", 0)

    # Add table of exercises/music
    doc.add_heading("Exercise - Music List", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False  # Disable autofit so manual widths are respected

    # Set column widths (adjust as needed)
    # column_widths = [Centimeters(0.7), Centimeters(6.0), Centimeters(9.0), Centimeters(2.5), Centimeters(5.0)]
    column_widths = [Cm(0.7), Cm(6.0), Cm(9.0), Cm(2.5), Cm(5.0)]
    for idx, width in enumerate(column_widths):
        table.columns[idx].width = width
    
    # Set header row
    table.style = 'Table Grid'
    

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Exercise"
    hdr_cells[2].text = "Music"
    hdr_cells[3].text = "Duration"
    hdr_cells[4].text = "Notes"

    all_songs = get_all_songs()
    # Convert all_songs to list of dicts if needed
    if all_songs and hasattr(all_songs[0], 'keys') and not isinstance(all_songs[0], dict):
        all_songs = [{k: s[k] for k in s.keys()} for s in all_songs]
    for idx, exercise_tuple in enumerate(session_exercises, 1):
        # Unpack tuple: (exercise_name, music_ref, exercise_id, notes)
        if len(exercise_tuple) >= 4:
            exercise_name, music_ref, exercise_id, notes = exercise_tuple
        else:
            exercise_name, music_ref, exercise_id = exercise_tuple
            notes = ""
        # Exercise column
        exercise_col = f"{exercise_name} "
        # Music details
        music_col = ""
        duration_col = ""
        if music_ref:
            song = next((s for s in all_songs if s["music_ref"] == music_ref), None)
            if song:
                music_col = f"{song['music_ref']} {song['title']} {{{song['artist']}}}"           
                duration_col = song.get("duration", "")
                # duration_parts = duration_col["duration"].split(":")
                # duration_col = f" 🕒 {int(duration_parts[0]) * 60 + int(duration_parts[1]):02}:{int(duration_parts[2]):02}"
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = exercise_col
        row_cells[2].text = music_col
        row_cells[3].text = duration_col
        row_cells[4].text = notes or ""

    
    # Add session metadata
    doc.add_heading("Session Metadata", level=1)
    for key, value in session_metadata.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.save(filepath)
    return filepath

if __name__ == "__main__":
    import argparse
    import pickle
    parser = argparse.ArgumentParser(description="Export session to MS Word.")
    parser.add_argument("--session_metadata", required=True, help="Path to pickled session_metadata dict")
    parser.add_argument("--session_exercises", required=True, help="Path to pickled session_exercises list")
    args = parser.parse_args()
    with open(args.session_metadata, "rb") as f:
        session_metadata = pickle.load(f)
    with open(args.session_exercises, "rb") as f:
        session_exercises = pickle.load(f)
    path = export_session_to_word(session_metadata, session_exercises)
    print(f"Exported to {path}")
